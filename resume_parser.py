"""Parse resumes from PDF, DOCX, or TXT into plain text."""

from __future__ import annotations

import os
from pathlib import Path

from pypdf import PdfReader
from docx import Document


def parse_resume(file_path: str) -> str:
    """Return plain text extracted from a resume file."""
    if not file_path or not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(file_path)
    if ext == ".docx":
        return _parse_docx(file_path)
    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    raise ValueError(f"Unsupported resume format: {ext}. Use .pdf, .docx, or .txt")


def _parse_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _parse_docx(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull text from tables (some resumes use them)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)
